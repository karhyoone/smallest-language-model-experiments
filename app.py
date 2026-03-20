import streamlit as st
from sentence_transformers import SentenceTransformer, util
import PyPDF2
import docx

st.set_page_config(page_title="AI Resume Matcher", layout="wide")

st.title("AI Resume Matcher & Job Analyzer")
st.markdown("Upload your resume + job description → get match score")

@st.cache_resource
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

model = load_model()

resume_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
job_file = st.file_uploader("Upload Job Description (PDF or DOCX)", type=["pdf", "docx"])
job_text_input = st.text_area("Or paste Job Description here", height=150)

if st.button("Analyze") and (resume_file or job_text_input):
    resume_text = ""
    if resume_file:
        if resume_file.type == "application/pdf":
            pdf = PyPDF2.PdfReader(resume_file)
            for page in pdf.pages:
                resume_text += page.extract_text() + "\n"
        else:
            doc = docx.Document(resume_file)
            resume_text = "\n".join(p.text for p in doc.paragraphs)

    job_text = job_text_input.strip()
    if not job_text and job_file:
        if job_file.type == "application/pdf":
            pdf = PyPDF2.PdfReader(job_file)
            for page in pdf.pages:
                job_text += page.extract_text() + "\n"
        else:
            doc = docx.Document(job_file)
            job_text = "\n".join(p.text for p in doc.paragraphs)

    if resume_text and job_text:
        resume_emb = model.encode(resume_text)
        job_emb = model.encode(job_text)
        similarity = util.cos_sim(resume_emb, job_emb)[0][0]
        score = round(float(similarity) * 100, 1)

        st.subheader(f"Match Score: {score}%")
        if score > 75:
            st.success("Strong match!")
        elif score > 50:
            st.info("Decent match - some improvements possible.")
        else:
            st.warning("Low match - add more relevant keywords/experience.")
    else:
        st.error("Need both resume and job text.")